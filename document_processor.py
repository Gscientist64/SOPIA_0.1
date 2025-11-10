# document_processor.py - Enhanced with Medical Knowledge Extraction
import os
import re
import math
import pickle
import time
from datetime import datetime
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

class DocumentProcessor:
    def __init__(self, sop_path, llm_client=None, checkpoint_file=None):
        self.sop_path = sop_path
        self.llm = llm_client
        self.checkpoint_file = checkpoint_file or f"{os.path.basename(sop_path)}_progress.pkl"
        self.progress = self._load_checkpoint()
        self.max_chunk_size = 4000
        self.chunk_overlap = 200
        self.retry_limit = 2
        self.retry_delay = 5

    def _load_checkpoint(self):
        if os.path.exists(self.checkpoint_file):
            print(f"[{self._now()}] 🔄 Found previous checkpoint — resuming progress.")
            with open(self.checkpoint_file, "rb") as f:
                return pickle.load(f)
        return {"completed_chunks": set(), "results": {}}

    def _save_checkpoint(self):
        with open(self.checkpoint_file, "wb") as f:
            pickle.dump(self.progress, f)
        print(f"[{self._now()}] 💾 Progress saved ({len(self.progress['completed_chunks'])} chunks done).")

    def _now(self):
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _extract_text(self, sop_path):
        """Extract ALL text content from DOCX including tables, headers, footers, etc."""
        doc = Document(sop_path)
        full_text = []
        
        print("🔍 Extracting document structure...")
        
        # Extract all paragraphs with their styles
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Include style information to detect headings
                style = para.style.name if para.style else "Normal"
                if style.lower() in ['heading 1', 'heading 2', 'heading 3', 'title']:
                    full_text.append(f"## {text} ##")  # Mark headings
                else:
                    full_text.append(text)
        
        # Extract tables with comprehensive structure
        table_count = 0
        for i, table in enumerate(doc.tables):
            table_count += 1
            table_text = [f"--- TABLE {i+1} START ---"]
            
            # Extract header row if it exists
            if len(table.rows) > 0:
                header_cells = []
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        header_cells.append(header_text)
                if header_cells:
                    table_text.append("HEADER: " + " | ".join(header_cells))
            
            # Extract all rows
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0 and any(cell.text.strip() for cell in row.cells):
                    continue  # Skip header if already processed
                    
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_text.append(f"Row {row_idx+1}: " + " | ".join(row_cells))
            
            table_text.append(f"--- TABLE {i+1} END ---")
            full_text.extend(table_text)
        
        print(f"✅ Extracted: {len(full_text)} elements, {table_count} tables")
        return "\n".join(full_text)

    def _create_granular_documents(self, text):
        """Create multiple document types from text to increase document count"""
        documents = []
        
        # Split into paragraphs and lines
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        for i, para in enumerate(paragraphs):
            # Skip very short lines that are probably formatting
            if len(para) < 20:
                continue
                
            # Create document for each paragraph
            documents.append({
                "text": para,
                "source_id": f"para_{i+1:04d}",
                "metadata": {"type": "paragraph"}
            })
            
            # Split long paragraphs into sentences for more granularity
            if len(para) > 100:
                sentences = re.split(r'[.!?]+', para)
                sentences = [s.strip() for s in sentences if len(s.strip()) > 30]
                
                for j, sentence in enumerate(sentences):
                    if len(sentence) > 30:  # Meaningful sentence
                        documents.append({
                            "text": sentence,
                            "source_id": f"para_{i+1:04d}_sent_{j+1}",
                            "metadata": {"type": "sentence"}
                        })
        
        # Extract key medical terms and create definition documents
        medical_terms = self._extract_medical_terms(text)
        for term, context in medical_terms:
            documents.append({
                "text": f"MEDICAL TERM: {term} - {context}",
                "source_id": f"term_{len(documents)+1:04d}",
                "metadata": {"type": "medical_term"}
            })
        
        # Extract table data as separate documents
        table_docs = self._extract_table_documents(text)
        documents.extend(table_docs)
        
        # Extract clinical algorithms and pathways
        pathway_docs = self._extract_clinical_pathways(text)
        documents.extend(pathway_docs)
        
        print(f"📄 Created {len(documents)} granular documents")
        return documents

    def _extract_medical_terms(self, text):
        """Extract medical terms and their context"""
        medical_keywords = [
            'ART', 'ARV', 'HTS', 'AHD', 'CD4', 'VL', 'viral load', 'TB', 
            'tuberculosis', 'prophylaxis', 'regimen', 'DTG', 'TDF', '3TC',
            'cryptococcal', 'meningitis', 'OI', 'opportunistic infection',
            'PEP', 'PrEP', 'IPT', 'STI', 'screening', 'diagnosis', 'treatment',
            'monitoring', 'adherence', 'counseling', 'testing', 'clinical stage',
            'WHO stage', 'immune reconstitution', 'IRIS', 'antiretroviral',
            'suppression', 'resistance', 'side effects', 'adverse events',
            'tenofovir', 'lamivudine', 'dolutegravir', 'efavirenz', 'nevirapine',
            'cotrimoxazole', 'isoniazid', 'preventive therapy', 'baseline',
            'follow-up', 'retention', 'viremia', 'suppressed', 'unsuppressed'
        ]
        
        terms_found = []
        lines = text.split('\n')
        
        for line in lines:
            for term in medical_keywords:
                if term.lower() in line.lower() and len(line) < 500:
                    terms_found.append((term, line.strip()))
                    # Don't break - allow multiple terms per line
                    
        # Remove duplicates while preserving order
        seen = set()
        unique_terms = []
        for term, context in terms_found:
            if (term, context) not in seen:
                seen.add((term, context))
                unique_terms.append((term, context))
        
        return unique_terms[:100]  # Limit to top 100 unique terms

    def _extract_table_documents(self, text):
        """Extract table data as separate documents"""
        table_docs = []
        lines = text.split('\n')
        
        current_table = None
        table_rows = []
        
        for line in lines:
            if 'TABLE' in line and 'START' in line:
                current_table = line
                table_rows = []
            elif 'TABLE' in line and 'END' in line:
                if current_table and table_rows:
                    # Create documents for each table row
                    for i, row in enumerate(table_rows):
                        table_docs.append({
                            "text": f"TABLE DATA: {row}",
                            "source_id": f"table_{len(table_docs)+1:04d}",
                            "metadata": {"type": "table_row"}
                        })
                current_table = None
                table_rows = []
            elif current_table and line.strip() and not line.startswith("---"):
                table_rows.append(line.strip())
        
        return table_docs

    def _extract_clinical_pathways(self, text):
        """Extract clinical pathways and algorithms"""
        pathway_docs = []
        lines = text.split('\n')
        
        # Look for algorithm descriptions
        algorithm_keywords = ['algorithm', 'pathway', 'flowchart', 'step', 'process', 'procedure']
        
        current_pathway = None
        pathway_steps = []
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in algorithm_keywords):
                # Start of a potential pathway
                if not current_pathway and len(line) > 20:
                    current_pathway = line
                    pathway_steps = [line]
                elif current_pathway:
                    pathway_steps.append(line)
            elif current_pathway and line.strip():
                pathway_steps.append(line)
            elif current_pathway and not line.strip():
                # End of pathway
                if len(pathway_steps) > 1:
                    pathway_text = "\n".join(pathway_steps)
                    pathway_docs.append({
                        "text": f"CLINICAL PATHWAY: {pathway_text}",
                        "source_id": f"pathway_{len(pathway_docs)+1:04d}",
                        "metadata": {"type": "clinical_pathway"}
                    })
                current_pathway = None
                pathway_steps = []
        
        return pathway_docs

    def _analyze_chunk(self, text, chunk_id):
        """Get comprehensive analysis focusing on medical and procedural content."""
        prompt = f"""
        Analyze this SOP section comprehensively. Extract ALL important information including:
        
        MEDICAL CONTENT:
        - Drug names, regimens, dosages
        - Medical procedures, protocols
        - Diagnostic criteria, lab values
        - Patient management guidelines
        - Definitions of medical terms (ART, ARV, HTS, AHD, etc.)
        
        PROCEDURAL CONTENT:
        - Step-by-step processes
        - Forms, documentation requirements
        - Timelines, schedules
        - Eligibility criteria
        - Monitoring requirements
        
        TABULAR DATA:
        - Extract all data from tables
        - Explain what each table shows
        - Note any important numbers, thresholds, criteria
        
        CLINICAL PATHWAYS:
        - Identify any clinical algorithms or decision trees
        - Extract if-then logic for clinical decisions
        - Note critical decision points
        
        CONTEXT: This is from an HIV Care Standard Operating Procedure document.
        
        SOP CONTENT:
        {text}
        
        Provide a structured analysis that captures EVERYTHING needed for accurate medical Q&A.
        Focus on preserving specific details, numbers, and exact medical terminology.
        """
        
        for attempt in range(1, self.retry_limit + 1):
            try:
                resp = self.llm.generate(prompt)
                if resp and len(resp.strip()) > 100:  # Ensure meaningful response
                    return resp.strip()
                else:
                    print(f"⚠️ Chunk {chunk_id} got short response, retrying...")
            except Exception as e:
                print(f"⚠️ Analysis failed for chunk {chunk_id} (attempt {attempt}): {e}")
                time.sleep(self.retry_delay * attempt)
        
        # Fallback: return key excerpts from the chunk
        print(f"⚠️ Using fallback analysis for chunk {chunk_id}")
        lines = text.split('\n')
        key_lines = [line for line in lines if any(marker in line for marker in 
                    ['TABLE', '##', 'ART', 'ARV', 'HTS', 'AHD', 'CD4', 'VL', 'prophylaxis', 'regimen', 'algorithm', 'step'])]
        return "KEY CONTENT: " + " | ".join(key_lines[:10]) if key_lines else "Content processed"

    def process(self):
        print(f"\n📘 Processing SOP: {self.sop_path}")
        start_time = time.time()
        
        # Extract text
        text = self._extract_text(self.sop_path)
        extraction_time = time.time() - start_time
        
        # Create granular documents (200-500+ documents)
        documents = self._create_granular_documents(text)
        
        # Add LLM analysis to some documents to enhance quality
        processed_docs = []
        analysis_count = 0
        
        for i, doc in enumerate(documents):
            # Add original document
            processed_docs.append(doc)
            
            # Add LLM analysis for important documents (every 10th document to manage API calls)
            if i % 10 == 0 and self.llm and len(doc['text']) > 100:
                try:
                    analysis = self._analyze_chunk(doc['text'], i+1)
                    if analysis and len(analysis) > 50:
                        processed_docs.append({
                            "text": f"ANALYSIS: {analysis}",
                            "source_id": f"analysis_{analysis_count+1:04d}",
                            "metadata": {"type": "analysis"}
                        })
                        analysis_count += 1
                        print(f"   → Added analysis {analysis_count} for document {i+1}")
                except Exception as e:
                    print(f"⚠️ Analysis failed for document {i+1}: {e}")
        
        processing_time = time.time() - start_time
        
        # Statistics
        doc_types = {}
        for doc in processed_docs:
            doc_type = doc['metadata']['type']
            doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
        
        print(f"🎯 Final document breakdown:")
        for doc_type, count in doc_types.items():
            print(f"   - {doc_type}: {count} documents")
        
        return {
            "title": os.path.splitext(os.path.basename(self.sop_path))[0],
            "documents": processed_docs,
            "chunks": processed_docs,  # Keep for compatibility
            "stats": {
                "total_documents": len(processed_docs),
                "document_types": doc_types,
                "analysis_count": analysis_count,
                "extraction_time": round(extraction_time, 1),
                "total_processing_time": round(processing_time, 1),
            }
        }

# Alias for compatibility
EnhancedDocumentProcessor = DocumentProcessor